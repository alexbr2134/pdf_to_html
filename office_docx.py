"""
Извлечение структуры DOCX: reading order (параграфы + таблицы) с merges.

Убирает layout-мусор: одноколоночные обёртки, пустые колонки/строки,
таблицы «Страница N», почти пустые декоративные сетки.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterator, Literal

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from office_table_regions import (
    find_trailing_chrome_start,
    grid_rows_as_prose_lines,
    looks_like_layout_form_table,
    should_emit_as_html_table,
    split_layout_table_grids,
    trim_trailing_form_chrome,
)
from pdf_html_pipeline import Cell


BlockKind = Literal["paragraph", "table", "heading"]


@dataclass
class DocxBlock:
    kind: BlockKind
    text: str = ""
    grid: list[list[Cell]] | None = None
    style: str | None = None
    heading_level: int | None = None


@dataclass
class DocxDocumentModel:
    blocks: list[DocxBlock] = field(default_factory=list)
    text: str = ""
    n_tables: int = 0
    n_images: int = 0
    n_embedded: int = 0
    total_cells: int = 0
    nonempty_cells: int = 0
    warnings: list[str] = field(default_factory=list)


_PAGE_ONLY_RE = re.compile(r"^\s*Страница\s+\d+\s*$", re.I)
_DASH_ONLY_RE = re.compile(r"^[\s\-–—_\.·•]+$")


def _cell_plain_text(cell) -> str:
    parts: list[str] = []
    for p in cell.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    # nested tables → append their text
    for nested in cell.tables:
        for row in nested.rows:
            for nc in row.cells:
                t = _cell_plain_text(nc)
                if t:
                    parts.append(t)
    return "\n".join(parts).strip()


def _tc_grid_span(tc) -> int:
    tcPr = tc.tcPr
    if tcPr is None or tcPr.gridSpan is None:
        return 1
    try:
        return max(1, int(tcPr.gridSpan.val))
    except Exception:
        return 1


def _tc_vmerge(tc) -> str | None:
    tcPr = tc.tcPr
    if tcPr is None or tcPr.vMerge is None:
        return None
    val = tcPr.vMerge.val
    if val is None or val == "continue":
        return "continue"
    return "restart"


def table_to_grid(table: Table) -> list[list[Cell]]:
    """Строит grid Cell с colspan/rowspan по OOXML (gridSpan / vMerge)."""
    rows_xml = table._tbl.tr_lst  # noqa: SLF001
    if not rows_xml:
        return []

    grid_cols = 0
    tbl_grid = table._tbl.tblGrid  # noqa: SLF001
    if tbl_grid is not None:
        grid_cols = len(tbl_grid.gridCol_lst)
    if grid_cols <= 0:
        for tr in rows_xml:
            span_sum = sum(_tc_grid_span(tc) for tc in tr.tc_lst)
            grid_cols = max(grid_cols, span_sum)
    grid_cols = max(grid_cols, 1)

    n_rows = len(rows_xml)
    slots: list[list[Any]] = [[None] * grid_cols for _ in range(n_rows)]

    for r_idx, tr in enumerate(rows_xml):
        c_idx = 0
        for tc in tr.tc_lst:
            while c_idx < grid_cols and slots[r_idx][c_idx] is not None:
                c_idx += 1
            if c_idx >= grid_cols:
                break

            colspan = _tc_grid_span(tc)
            vmerge = _tc_vmerge(tc)
            text = ""
            try:
                for cell in table.rows[r_idx].cells:
                    if cell._tc is tc:  # noqa: SLF001
                        text = _cell_plain_text(cell)
                        break
            except Exception:
                texts = [(node.text or "") for node in tc.iter(qn("w:t"))]
                text = " ".join(t.strip() for t in texts if t and t.strip())

            if vmerge == "continue":
                up = r_idx - 1
                while up >= 0 and slots[up][c_idx] == "cov":
                    up -= 1
                master = slots[up][c_idx] if up >= 0 else None
                if isinstance(master, Cell):
                    master.rowspan = max(master.rowspan, r_idx - master.row + 1)
                for dc in range(colspan):
                    if c_idx + dc < grid_cols:
                        slots[r_idx][c_idx + dc] = "cov"
                c_idx += colspan
                continue

            cell = Cell(
                row=r_idx,
                col=c_idx,
                bbox=None,
                text=text,
                rowspan=1,
                colspan=colspan,
                covered=False,
            )
            slots[r_idx][c_idx] = cell
            for dc in range(1, colspan):
                if c_idx + dc < grid_cols:
                    slots[r_idx][c_idx + dc] = "cov"
            c_idx += colspan

    grid: list[list[Cell]] = []
    for r_idx in range(n_rows):
        row_cells: list[Cell] = []
        for c_idx in range(grid_cols):
            slot = slots[r_idx][c_idx]
            if isinstance(slot, Cell):
                row_cells.append(slot)
            elif slot == "cov":
                row_cells.append(
                    Cell(
                        row=r_idx,
                        col=c_idx,
                        bbox=None,
                        text="",
                        covered=True,
                        is_placeholder=True,
                    )
                )
            else:
                row_cells.append(
                    Cell(
                        row=r_idx,
                        col=c_idx,
                        bbox=None,
                        text="",
                        is_placeholder=True,
                    )
                )
        grid.append(row_cells)
    return grid


def compact_grid(grid: list[list[Cell]]) -> list[list[Cell]]:
    """Удаляет пустые строки/колонки, поджимает colspan."""
    if not grid:
        return grid
    n_rows = len(grid)
    n_cols = len(grid[0])

    keep_cols: list[int] = []
    for c in range(n_cols):
        useful = False
        for r in range(n_rows):
            cell = grid[r][c]
            if cell.covered:
                continue
            if (cell.text or "").strip() or cell.colspan > 1 or cell.rowspan > 1:
                useful = True
                break
        if useful:
            keep_cols.append(c)
    if not keep_cols:
        return []

    col_map = {old: i for i, old in enumerate(keep_cols)}
    keep_set = set(keep_cols)
    out: list[list[Cell]] = []

    for row in grid:
        has = False
        for c in keep_cols:
            cell = row[c]
            if cell.covered:
                continue
            if (cell.text or "").strip() or cell.colspan > 1:
                has = True
                break
        if not has:
            continue
        new_row: list[Cell] = []
        for c in keep_cols:
            cell = row[c]
            if cell.covered:
                new_row.append(
                    Cell(
                        row=len(out),
                        col=col_map[c],
                        bbox=None,
                        text="",
                        covered=True,
                        is_placeholder=True,
                    )
                )
                continue
            cs = cell.colspan
            if cs > 1:
                spanned = [c + k for k in range(cs) if c + k < n_cols]
                kept = [x for x in spanned if x in keep_set]
                cs = max(1, len(kept))
            text = (cell.text or "").strip()
            # drop dash-only noise
            if text and _DASH_ONLY_RE.match(text) and len(text) < 4:
                text = ""
            new_row.append(
                Cell(
                    row=len(out),
                    col=col_map[c],
                    bbox=None,
                    text=text,
                    rowspan=min(cell.rowspan, max(1, n_rows)),
                    colspan=cs,
                    covered=False,
                    is_placeholder=not bool(text),
                )
            )
        out.append(new_row)

    for r, row in enumerate(out):
        for cell in row:
            if not cell.covered and cell.rowspan > 1:
                cell.rowspan = min(cell.rowspan, len(out) - r)
    return out


def _nonempty_stats(grid: list[list[Cell]]) -> tuple[int, int, int]:
    """visible_cells, nonempty, max_text_col_count."""
    visible = 0
    nonempty = 0
    cols_with_text = set()
    for row in grid:
        for cell in row:
            if cell.covered:
                continue
            visible += 1
            if (cell.text or "").strip():
                nonempty += 1
                cols_with_text.add(cell.col)
    return visible, nonempty, len(cols_with_text)


def classify_table_block(grid: list[list[Cell]]) -> Literal["table", "prose", "drop"]:
    """table / prose / drop — оставляем только нормальные data-grid."""
    if not grid:
        return "drop"
    grid = compact_grid(grid)
    if not grid:
        return "drop"

    n_rows = len(grid)
    n_cols = len(grid[0])
    visible, nonempty, text_cols = _nonempty_stats(grid)

    texts = [
        (c.text or "").strip()
        for row in grid
        for c in row
        if not c.covered and (c.text or "").strip()
    ]
    if not texts:
        return "drop"

    if all(_PAGE_ONLY_RE.match(t) for t in texts):
        return "drop"

    if n_rows == 1 and n_cols == 1:
        return "prose"
    if text_cols <= 1 and (nonempty <= 12 or n_cols == 1):
        return "prose"

    if visible >= 8 and nonempty / max(1, visible) < 0.12 and text_cols <= 2:
        return "prose"

    if should_emit_as_html_table(grid):
        return "table"
    return "prose"


def grid_to_prose_blocks(grid: list[list[Cell]]) -> list[DocxBlock]:
    return [
        DocxBlock(kind="paragraph", text=line)
        for line in grid_rows_as_prose_lines(grid)
        if line.strip()
    ]


def _iter_block_items(parent: DocumentObject) -> Iterator[Paragraph | Table]:
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _heading_level(paragraph: Paragraph) -> int | None:
    style = paragraph.style
    if style is None:
        return None
    name = (style.name or "").strip()
    m = re.match(r"Heading\s*(\d+)$", name, re.I)
    if m:
        return int(m.group(1))
    m = re.match(r"Заголовок\s*(\d+)$", name, re.I)
    if m:
        return int(m.group(1))
    return None


def _count_drawings(document: DocumentObject) -> tuple[int, int]:
    body = document.element.body
    n_drawing = len(list(body.iter(qn("w:drawing"))))
    n_object = len(list(body.iter(qn("w:object"))))
    n_pict = len(list(body.iter(qn("w:pict"))))
    return n_drawing + n_pict, n_object


def extract_docx(path: str | Path) -> DocxDocumentModel:
    """Парсит DOCX в упорядоченные блоки с очисткой layout-таблиц."""
    src = Path(path)
    try:
        document = Document(str(src))
    except Exception as exc:
        msg = str(exc).lower()
        if "password" in msg or "encrypt" in msg:
            raise PermissionError("encrypted") from exc
        raise

    blocks: list[DocxBlock] = []
    text_parts: list[str] = []
    n_tables = 0
    total_cells = 0
    nonempty_cells = 0
    warnings: list[str] = []

    for item in _iter_block_items(document):
        if isinstance(item, Paragraph):
            t = (item.text or "").strip()
            if not t or _DASH_ONLY_RE.match(t):
                continue
            level = _heading_level(item)
            kind: BlockKind = "heading" if level else "paragraph"
            blocks.append(
                DocxBlock(
                    kind=kind,
                    text=t,
                    style=(item.style.name if item.style else None),
                    heading_level=level,
                )
            )
            text_parts.append(t)
            continue

        raw = table_to_grid(item)
        grid = compact_grid(raw)
        if not grid:
            warnings.append("empty table dropped")
            continue

        # форма целиком → сначала режем на пояса, потом решаем table/prose
        grids: list[list[list[Cell]]] = [grid]
        if looks_like_layout_form_table(grid):
            split = split_layout_table_grids(
                grid, table=item, raw_row_count=len(raw)
            )
            if len(split) > 1:
                grids = split
                warnings.append(f"split→{len(split)}")

        for g in grids:
            chrome_at = find_trailing_chrome_start(g)
            tail = None
            if chrome_at is not None and chrome_at >= 3:
                tail = g[chrome_at:]
                g = g[:chrome_at]
            g = trim_trailing_form_chrome(g)
            kind = classify_table_block(g)
            if kind == "drop":
                if tail:
                    prose = grid_to_prose_blocks(tail)
                    blocks.extend(prose)
                    for b in prose:
                        text_parts.append(b.text)
                continue
            if kind == "prose":
                prose = grid_to_prose_blocks(g)
                blocks.extend(prose)
                for b in prose:
                    text_parts.append(b.text)
            else:
                n_tables += 1
                cell_texts: list[str] = []
                for row in g:
                    for cell in row:
                        if cell.covered:
                            continue
                        total_cells += 1
                        if (cell.text or "").strip():
                            nonempty_cells += 1
                            cell_texts.append(cell.text.strip())
                blocks.append(
                    DocxBlock(kind="table", grid=g, text="\n".join(cell_texts))
                )
                text_parts.extend(cell_texts)
            if tail:
                prose = grid_to_prose_blocks(tail)
                blocks.extend(prose)
                for b in prose:
                    text_parts.append(b.text)

    n_images, n_embedded = _count_drawings(document)
    return DocxDocumentModel(
        blocks=blocks,
        text="\n".join(text_parts),
        n_tables=n_tables,
        n_images=n_images,
        n_embedded=n_embedded,
        total_cells=total_cells,
        nonempty_cells=nonempty_cells,
        warnings=warnings,
    )
